from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import re
import os
import time
import platform

# ไลบรารีสำหรับควบคุม Word COM (Windows เท่านั้น)
import win32com.client
import pythoncom

from Backend import Data

# ตั้งค่า font ภาษาไทย
def set_font_thai(run, size_pt=15, bold=False):
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run.font.size = Pt(size_pt)
    run.font.bold = bold

# แปลงข้อความจาก textarea พร้อมจัดการย่อหน้าตามช่องว่างหรือ tab
def prepare_body_paragraphs(doc, raw_text):
    # Normalize newlines to ensure consistency across OS
    lines = raw_text.replace('\r\n', '\n').split('\n')
    
    # List to accumulate lines for the current logical paragraph
    current_paragraph_lines = []

    # Function to add a logical paragraph to the document
    def add_logical_paragraph_to_doc():
        nonlocal current_paragraph_lines # Allow modifying the outer scope variable
        if current_paragraph_lines: # Only add if there is content
            para = doc.add_paragraph()
            # Apply first line indent for all generated content paragraphs
            para.paragraph_format.first_line_indent = Cm(1.27)
            # Join all lines collected for this logical paragraph, removing extra spaces
            run = para.add_run(" ".join(line.strip() for line in current_paragraph_lines if line.strip()))
            set_font_thai(run, size_pt=15)
            # Add some space after a non-empty paragraph for better readability
            para.paragraph_format.space_after = Pt(6)
            current_paragraph_lines = [] # Reset for the next paragraph

    for i, line in enumerate(lines):
        # Check if the line starts with 14 or more spaces OR a tab character.
        # This check must be performed on the original line BEFORE stripping,
        # as stripping removes the very indentation we are looking for.
        is_explicitly_indented_start = False
        if len(line) >= 14 and line[:14].isspace():
            is_explicitly_indented_start = True
        elif line.startswith('\t\t'):
            is_explicitly_indented_start = True
        
        stripped_line = line.strip()

        if stripped_line == "":
            # If a blank line is encountered, finalize the current logical paragraph
            # and then add an empty paragraph to represent the blank line in Word.
            add_logical_paragraph_to_doc() # Finalize previous paragraph
            
            # Add an empty paragraph for the blank line itself
            blank_para = doc.add_paragraph()
            set_font_thai(blank_para.add_run(""), size_pt=15) # Apply font even if empty
            blank_para.paragraph_format.space_before = Pt(0) # No extra space before/after blank line
            blank_para.paragraph_format.space_after = Pt(0)
        elif is_explicitly_indented_start:
            # If this line explicitly triggers a new indented paragraph,
            # finalize the previous accumulated paragraph (if any).
            add_logical_paragraph_to_doc() # Finalize previous paragraph
            
            # Start a new paragraph with the content of the current line (after stripping)
            current_paragraph_lines.append(stripped_line)
        else:
            # Otherwise, it's a continuation of the current logical paragraph
            # or the very first line of the body text (if not explicitly indented).
            current_paragraph_lines.append(stripped_line)

    # After the loop, add any remaining accumulated content as the last paragraph
    add_logical_paragraph_to_doc()

# ปิดเฉพาะแท็บไฟล์ Word ที่เปิดอยู่ (ถ้ามี)
def close_word_file_if_open(filename):
    pythoncom.CoInitialize()  # เรียก COM สำหรับ thread นี้
    try:
        word = win32com.client.Dispatch("Word.Application")
        for doc in word.Documents:
            # เปรียบเทียบชื่อไฟล์แบบไม่สนใจ case
            if filename.lower() in doc.FullName.lower():
                print(f"📄 พบไฟล์ {filename} ที่เปิดอยู่ใน Word — กำลังปิดแท็บ")
                doc.Close(False)  # False = ปิดโดยไม่บันทึกซ้ำ
                return True
    except Exception as e:
        print("❌ ไม่สามารถตรวจสอบหรือปิดเอกสาร Word:", e)
    return False

# บันทึกไฟล์พร้อม retry และปิดแท็บ Word เฉพาะไฟล์นั้นถ้ายังเปิดอยู่
def save_doc_with_retry(doc, filename="Sleeve1_Output.docx", max_retries=3):
    for attempt in range(max_retries):
        try:
            doc.save(filename)
            print(f"✅ {filename} created successfully!")
            if platform.system() == "Windows":
                os.startfile(filename)
            return True
        except PermissionError:
            print(f"⚠️ ไม่สามารถบันทึกไฟล์ {filename} ได้ อาจยังเปิดอยู่ใน Word")
            print("🔄 กำลังพยายามปิดเฉพาะแท็บของไฟล์นั้น...")
            closed = close_word_file_if_open(filename)
            if not closed:
                print("⏳ รอ 2 วินาทีแล้วลองใหม่...")
            time.sleep(2)
    print("❌ ไม่สามารถบันทึกไฟล์ได้ กรุณาปิดไฟล์ด้วยตนเองแล้วลองใหม่อีกครั้ง")
    return False

def Sleeve1(Data, title, runNumber, bodyText1):
    doc = Document()

    # ตั้งค่าสไตล์ปกติ (Normal style) สำหรับฟอนต์ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(15)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ตั้งค่าระยะขอบหน้ากระดาษ
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # เพิ่มโลโก้ที่มุมซ้ายบน
    section.different_first_page_header_footer = True

    try:
        header_first = section.first_page_header
        header_first_paragraph = header_first.paragraphs[0] if header_first.paragraphs else header_first.add_paragraph()
        run = header_first_paragraph.add_run()
        run.add_picture("Logo.jpg", width=Cm(1.91))  # ปรับขนาดตามต้องการ
    except Exception as e:
        print("⚠️ ไม่สามารถแทรกรูปโลโก้ได้:", e)

    # หัวข้อ "บันทึกข้อความ"
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("บันทึกข้อความ")
    set_font_thai(run, size_pt=22, bold=True)

    # ส่วนราชการ
    p_gov_section = doc.add_paragraph("ส่วนราชการ ภาควิชาอุตสาหกรรมเกษตร คณะเกษตรศาสตร์ฯ ทรัพยากรธรรมชาติและสิ่งแวดล้อม โทร. 2749")
    set_font_thai(p_gov_section.runs[0], size_pt=15)
    p_gov_section.paragraph_format.space_after = Pt(0)

    # ที่ และ วันที่
    p_ref_date = doc.add_paragraph()
    run_ref = p_ref_date.add_run(f"ที่ {runNumber}")
    set_font_thai(run_ref, size_pt=15)
    p_ref_date.paragraph_format.space_after = Pt(0)
    p_ref_date.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_PARAGRAPH_ALIGNMENT.RIGHT)
    p_ref_date.add_run("\t")
    date_run = p_ref_date.add_run(f"วันที่ {Data.day}")
    set_font_thai(date_run, size_pt=15)

    # เรื่อง
    p_subject = doc.add_paragraph(f"เรื่อง {title}")
    set_font_thai(p_subject.runs[0], size_pt=15)
    p_subject.paragraph_format.space_after = Pt(0)

    # เส้นคั่น
    p_line = doc.add_paragraph()
    run_line = p_line.add_run("-" * 139)
    set_font_thai(run_line, size_pt=15)
    p_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_line.paragraph_format.space_after = Pt(0)
    p_line.paragraph_format.space_before = Pt(0)

    # เรียน
    p_dean = doc.add_paragraph("เรียน คณบดีคณะเกษตรศาสตร์ฯ")
    set_font_thai(p_dean.runs[0], size_pt=15)
    p_dean.paragraph_format.space_before = Pt(0)
    p_dean.paragraph_format.space_after = Pt(12)

    # เนื้อความหลักของบันทึกข้อความ
    prepare_body_paragraphs(doc, bodyText1)

    # ตารางรายการ
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    headers = [
        "ลำดับ",
        "รายการขอซื้อจ้าง\n[ขั้นตอนตามระเบียบฯ ข้อ 22(2)]",
        "รายการแยกวัสดุ\nตามระบบบัญชี\n3 มิติ",
        "จำนวนหน่วย",
        "กำหนดเวลาที่\nต้องการใช้พัสดุ\n[ตามระเบียบข้อ\n22 (5)]"
    ]

    for i, header_text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header_text)
        set_font_thai(run, bold=True, size_pt=14)
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, row_data in enumerate(Data.list, start=1):
        cells = table.add_row().cells

        item_name = row_data[0] if len(row_data) > 0 else ""
        category = row_data[1] if len(row_data) > 1 else ""
        quantity = row_data[2] if len(row_data) > 2 else ""
        date_needed = row_data[3] if len(row_data) > 3 else ""

        run_idx = cells[0].paragraphs[0].add_run(str(idx))
        set_font_thai(run_idx, size_pt=15)
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        run_item = cells[1].paragraphs[0].add_run(item_name)
        set_font_thai(run_item, size_pt=15)
        cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        run_category = cells[2].paragraphs[0].add_run(category)
        set_font_thai(run_category, size_pt=15)
        cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        run_quantity = cells[3].paragraphs[0].add_run(quantity)
        set_font_thai(run_quantity, size_pt=15)
        cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[3].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        run_date_needed = cells[4].paragraphs[0].add_run(date_needed)
        set_font_thai(run_date_needed, size_pt=15)
        cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells[4].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for _ in range(3):
        doc.add_paragraph()

    # ลงชื่อ
    p_signature = doc.add_paragraph()
    p_signature.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_sig = p_signature.add_run("ลงชื่อ ..........................................................")
    set_font_thai(run_sig, size_pt=15)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_name = p_name.add_run("(รศ.ดร.ทิพวรรณ ทองสุข)")
    set_font_thai(run_name, size_pt=15)

    return save_doc_with_retry(doc)

# ทดสอบ (หากรันไฟล์นี้โดยตรง)
if __name__ == '__main__':
        # สร้าง instance ของคลาส Data
    data_handler = Data()

    # --- ข้อมูลตัวอย่างที่มี บริษัทซ้ำ วันซ้ำ แต่คนละใบรับ ---

    # วันที่ 10/มิ.ย./68
    date_a = "10 มิ.ย.68"
    company_a = "บริษัท สยามพัฒนา จำกัด"
    company_b = "ร้านค้าส่งอุปกรณ์"

    # ใบรับที่ INV_A_001
    data_handler.appendlist("ปากกาเคมี", "วัสดุสำนักงาน", "5 ด้าม", date_a, 20.00, company_a, "INV_A_001", date_a)
    data_handler.appendlist("กระดาษโน้ต", "วัสดุสำนักงาน", "3 เล่ม", date_a, 15.00, company_a, "INV_A_001", date_a)
    
    # ใบรับที่ INV_A_002 (บริษัทเดียวกัน วันเดียวกัน แต่คนละใบรับ)
    data_handler.appendlist("แฟ้มเอกสาร", "วัสดุสำนักงาน", "10 อัน", date_a, 25.00, company_a, "INV_A_002", date_a)
    data_handler.appendlist("ยางลบ", "วัสดุสำนักงาน", "5 ก้อน", date_a, 5.00, company_a, "INV_A_002", date_a)

    # ใบรับที่ INV_B_001
    data_handler.appendlist("เทปใส", "วัสดุสำนักงาน", "2 ม้วน", date_a, 10.00, company_b, "INV_B_001", date_a)
    
    # ใบรับที่ INV_B_002 (บริษัทเดียวกัน วันเดียวกัน แต่คนละใบรับ)
    data_handler.appendlist("กาวแท่ง", "วัสดุสำนักงาน", "5 แท่ง", date_a, 8.00, company_b, "INV_B_002", date_a)
    data_handler.appendlist("ไม้บรรทัด", "วัสดุสำนักงาน", "3 อัน", date_a, 12.00, company_b, "INV_B_002", date_a)


    # วันที่ 15/มิ.ย./68
    date_b = "15 มิ.ย. 68"
    company_c = "บริษัท วัสดุก่อสร้าง"
    company_d = "บริษัท สยามพัฒนา จำกัด" # บริษัทซ้ำจาก date_a แต่คนละวัน คนละใบรับ

    data_handler.appendlist("หลอดไฟ LED", "วัสดุไฟฟ้า", "10 หลอด", date_b, 45.00, company_c, "INV_C_001", date_b)
    data_handler.appendlist("สายไฟ", "วัสดุไฟฟ้า", "1 ม้วน", date_b, 200.00, company_c, "INV_C_001", date_b)
    data_handler.appendlist("เม้าส์ไร้สาย", "อุปกรณ์คอมพิวเตอร์", "1 ชิ้น", date_b, 350.00, company_d, "INV_D_001", date_b)

    # วันที่ 20/มิ.ย./68
    date_c = "20 มิ.ย. 68"
    company_e = "ร้านค้าส่งอุปกรณ์" # บริษัทซ้ำจาก date_a แต่คนละวัน คนละใบรับ
    company_f = "บริษัท อาหารสด"

    data_handler.appendlist("ปากกาเจล", "วัสดุสำนักงาน", "12 ด้าม", date_c, 18.00, company_e, "INV_E_001", date_c)
    data_handler.appendlist("นมสด", "วัสดุบริโภค", "6 กล่อง", date_c, 30.00, company_f, "INV_F_001", date_c)

    data_handler.appendlist("เม้าส์ไร้สาย", "อุปกรณ์คอมพิวเตอร์", "1 ชิ้น", date_b, 350.00, company_d, "INV_D_099", date_b)

    example_text = data_handler.generate_purchase_request("a","a","a","a","a")
    Sleeve1(data_handler, "a", "a", example_text)
